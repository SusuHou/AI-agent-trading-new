"""Build the English methodology working draft as a Word document."""

from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_PATH = Path(__file__).resolve().parents[1] / "docs" / "methodology" / (
    "Methodology_Working_Draft_Steps_01_to_14.docx"
)

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
NAVY = RGBColor(11, 37, 69)
GRAY = RGBColor(89, 96, 105)
LIGHT_GRAY_HEX = "F4F6F9"


def set_font(run, name: str, size: float | None = None) -> None:
    """Set a font in a way that Word and LibreOffice both honor."""

    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)


def set_style_font(style, name: str, size: float, color: RGBColor | None = None) -> None:
    """Apply explicit font tokens to a paragraph style."""

    style.font.name = name
    style.font.size = Pt(size)
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if color is not None:
        style.font.color.rgb = color


def set_paragraph_shading(paragraph, fill: str) -> None:
    """Add a restrained background fill to a paragraph."""

    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def add_page_number(paragraph) -> None:
    """Insert a live Word PAGE field."""

    run = paragraph.add_run("Page ")
    set_font(run, "Calibri", 9)
    run.font.color.rgb = GRAY

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    field_run = OxmlElement("w:r")
    field_run.append(begin)
    field_run.append(instruction)
    field_run.append(separate)
    field_run.append(text)
    field_run.append(end)
    paragraph._p.append(field_run)


def add_body_paragraph(doc: Document, text: str):
    """Add a body paragraph using the document's narrative preset."""

    replacements = [
        ("E[v_t | y_t]", "E[vₜ | yₜ]"),
        ("sum_i x_i,t", "Σᵢ xᵢ,ₜ"),
        ("sigma_v_hat", "σ̂ᵥ"),
        ("lambda^N", "λᴺ"),
        ("lambda^M", "λᴹ"),
        ("chi^N", "χᴺ"),
        ("chi^M", "χᴹ"),
        ("xi^2", "ξ²"),
        ("x_1,t", "x₁,ₜ"),
        ("x_2,t", "x₂,ₜ"),
        ("x_i,t", "xᵢ,ₜ"),
        ("sigma_v", "σᵥ"),
        ("sigma_u", "σᵤ"),
        ("n_x", "nₓ"),
        ("n_p", "nₚ"),
        ("n_v", "nᵥ"),
        ("iota", "ι"),
        ("v-bar", "v̄"),
        ("v_t", "vₜ"),
        ("p_t", "pₜ"),
        ("y_t", "yₜ"),
        ("z_t", "zₜ"),
        ("u_t", "uₜ"),
    ]
    for old, new in replacements:
        text = text.replace(old, new)
    text = re.sub(r"\bxi\b", "ξ", text)
    text = re.sub(r"\btheta\b", "θ", text)
    text = re.sub(r"\blambda\b", "λ", text)
    text = re.sub(r"\bchi\b", "χ", text)
    text = re.sub(r"\bgamma\b", "γ", text)

    paragraph = doc.add_paragraph(text, style="Normal")
    paragraph.paragraph_format.widow_control = True
    return paragraph


def add_equation(doc: Document, text: str):
    """Add a centered, editable equation line."""

    if doc.paragraphs:
        doc.paragraphs[-1].paragraph_format.keep_with_next = True
    paragraph = doc.add_paragraph(style="Equation")
    run = paragraph.add_run(text)
    set_font(run, "Cambria Math", 11.5)
    return paragraph


def add_draft_note(doc: Document, label: str, text: str) -> None:
    """Add a one-paragraph scope note without using a layout table."""

    paragraph = doc.add_paragraph(style="Draft Note")
    label_run = paragraph.add_run(f"{label} ")
    label_run.bold = True
    set_font(label_run, "Calibri", 10.5)
    text_run = paragraph.add_run(text)
    set_font(text_run, "Calibri", 10.5)
    set_paragraph_shading(paragraph, LIGHT_GRAY_HEX)


def set_cell_shading(cell, fill: str) -> None:
    """Apply a solid background fill to one table cell."""

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top: int, start: int, bottom: int, end: int) -> None:
    """Set table-cell padding in DXA units."""

    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for edge, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, column_widths: list[int]) -> None:
    """Apply deterministic table width, indent, grid, and cell widths."""

    total_width = sum(column_widths)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_width))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_grid = table._tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in column_widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for cell, width in zip(row.cells, column_widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell, top=100, start=120, bottom=100, end=120)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_fixed_point_results_table(doc: Document) -> None:
    """Add the four calibrated benchmark coefficient rows."""

    caption = doc.add_paragraph(
        "Table 1. Fixed-point coefficients under the two paper noise environments",
        style="Caption",
    )
    caption.paragraph_format.keep_with_next = True

    headers = ("Noise σᵤ", "Benchmark", "λ", "γ", "χ")
    rows = (
        ("Low (0.1)", "Nash", "0.002000000400", "0.003000000293", "166.666633333"),
        ("Low (0.1)", "Perfect cartel", "0.002000000800", "0.004000000873", "124.999950000"),
        ("High (100)", "Nash", "0.002000000289", "0.002721588765", "166.666642614"),
        ("High (100)", "Perfect cartel", "0.002000000554", "0.003384490171", "124.999965388"),
    )

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, [1500, 1800, 2020, 2020, 2020])

    header_row = table.rows[0]
    tr_pr = header_row._tr.get_or_add_trPr()
    repeat_header = OxmlElement("w:tblHeader")
    repeat_header.set(qn("w:val"), "true")
    tr_pr.append(repeat_header)

    for cell, text in zip(header_row.cells, headers):
        set_cell_shading(cell, "D9EAF7")
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(text)
        set_font(run, "Calibri", 9.5)
        run.bold = True
        run.font.color.rgb = NAVY

    for row_values in rows:
        row = table.add_row()
        for cell, text in zip(row.cells, row_values):
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.05
            run = paragraph.add_run(text)
            set_font(run, "Calibri", 9.2)

    set_table_geometry(table, [1500, 1800, 2020, 2020, 2020])


def add_grid_summary_table(doc: Document) -> None:
    """Compare the validated action- and price-grid ranges."""

    caption = doc.add_paragraph(
        "Table 2. Validated finite-grid ranges under the two noise environments",
        style="Caption",
    )
    caption.paragraph_format.keep_with_next = True

    headers = (
        "Noise σᵤ",
        "Action multiplier interval [cᴸ, cᴴ]",
        "Final price interval",
        "Price spacing",
    )
    rows = (
        (
            "Low (0.1)",
            "[120.833281667, 170.833301667]",
            "[-0.316353302, 2.316353302]",
            "0.087756887",
        ),
        (
            "High (100)",
            "[120.833297665, 170.833310336]",
            "[-0.786282969, 2.786282969]",
            "0.119085531",
        ),
    )

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, [1250, 3150, 2900, 2060])

    header_row = table.rows[0]
    tr_pr = header_row._tr.get_or_add_trPr()
    repeat_header = OxmlElement("w:tblHeader")
    repeat_header.set(qn("w:val"), "true")
    tr_pr.append(repeat_header)

    for cell, text in zip(header_row.cells, headers):
        set_cell_shading(cell, "D9EAF7")
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(text)
        set_font(run, "Calibri", 9.2)
        run.bold = True
        run.font.color.rgb = NAVY

    for row_values in rows:
        row = table.add_row()
        for cell, text in zip(row.cells, row_values):
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            run = paragraph.add_run(text)
            set_font(run, "Calibri", 9.0)

    set_table_geometry(table, [1250, 3150, 2900, 2060])


def configure_styles(doc: Document) -> None:
    """Resolve and apply the narrative_proposal design preset."""

    styles = doc.styles

    normal = styles["Normal"]
    set_style_font(normal, "Calibri", 11, RGBColor(0, 0, 0))
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.widow_control = True

    title = styles["Title"]
    set_style_font(title, "Calibri", 24, NAVY)
    title.font.bold = True
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(5)
    title.paragraph_format.keep_with_next = True
    title_p_pr = title._element.get_or_add_pPr()
    for border in title_p_pr.findall(qn("w:pBdr")):
        title_p_pr.remove(border)

    subtitle = styles["Subtitle"]
    set_style_font(subtitle, "Calibri", 12.5, GRAY)
    subtitle.font.italic = False
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(15)
    subtitle.paragraph_format.keep_with_next = True

    heading_1 = styles["Heading 1"]
    set_style_font(heading_1, "Calibri", 16, BLUE)
    heading_1.font.bold = True
    heading_1.paragraph_format.space_before = Pt(18)
    heading_1.paragraph_format.space_after = Pt(10)
    heading_1.paragraph_format.keep_with_next = True
    heading_1.paragraph_format.keep_together = True

    heading_2 = styles["Heading 2"]
    set_style_font(heading_2, "Calibri", 13, BLUE)
    heading_2.font.bold = True
    heading_2.paragraph_format.space_before = Pt(12)
    heading_2.paragraph_format.space_after = Pt(6)
    heading_2.paragraph_format.keep_with_next = True
    heading_2.paragraph_format.keep_together = True

    heading_3 = styles["Heading 3"]
    set_style_font(heading_3, "Calibri", 12, DARK_BLUE)
    heading_3.font.bold = True
    heading_3.paragraph_format.space_before = Pt(8)
    heading_3.paragraph_format.space_after = Pt(4)
    heading_3.paragraph_format.keep_with_next = True
    heading_3.paragraph_format.keep_together = True

    equation = styles.add_style("Equation", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(equation, "Cambria Math", 11.5, RGBColor(0, 0, 0))
    equation.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    equation.paragraph_format.left_indent = Inches(0.25)
    equation.paragraph_format.right_indent = Inches(0.25)
    equation.paragraph_format.space_before = Pt(5)
    equation.paragraph_format.space_after = Pt(8)
    equation.paragraph_format.line_spacing = 1.0
    equation.paragraph_format.keep_together = True
    equation.paragraph_format.keep_with_next = True
    equation.paragraph_format.widow_control = True

    note = styles.add_style("Draft Note", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(note, "Calibri", 10.5, RGBColor(44, 52, 60))
    note.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    note.paragraph_format.left_indent = Inches(0.15)
    note.paragraph_format.right_indent = Inches(0.15)
    note.paragraph_format.space_before = Pt(4)
    note.paragraph_format.space_after = Pt(12)
    note.paragraph_format.line_spacing = 1.2
    note.paragraph_format.keep_together = True

    reference = styles.add_style("Reference", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(reference, "Calibri", 10.5, RGBColor(0, 0, 0))
    reference.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    reference.paragraph_format.left_indent = Inches(0.3)
    reference.paragraph_format.first_line_indent = Inches(-0.3)
    reference.paragraph_format.space_before = Pt(0)
    reference.paragraph_format.space_after = Pt(6)
    reference.paragraph_format.line_spacing = 1.2

    caption = styles["Caption"]
    set_style_font(caption, "Calibri", 9.5, DARK_BLUE)
    caption.font.bold = True
    caption.font.italic = False
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    caption.paragraph_format.space_before = Pt(8)
    caption.paragraph_format.space_after = Pt(5)
    caption.paragraph_format.keep_with_next = True
    caption.paragraph_format.keep_together = True


def configure_page(doc: Document) -> None:
    """Apply the preset's exact US Letter page geometry and quiet furniture."""

    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1.2)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    # Academic-manuscript override: use no running header. Populate all three
    # footer stories explicitly so page numbering remains stable across Word
    # applications with different first/odd/even page settings.
    doc.settings.odd_and_even_pages_header_footer = True
    section.different_first_page_header_footer = True

    for footer in (
        section.footer,
        section.even_page_footer,
        section.first_page_footer,
    ):
        footer_paragraph = footer.paragraphs[0]
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer_paragraph.paragraph_format.space_before = Pt(0)
        footer_paragraph.paragraph_format.space_after = Pt(0)
        add_page_number(footer_paragraph)


def build_document() -> Document:
    """Create the methodology working draft."""

    doc = Document()
    configure_styles(doc)
    configure_page(doc)

    properties = doc.core_properties
    properties.title = (
        "Methodology Working Draft: Market Design, Benchmarks, and Finite States"
    )
    properties.subject = "Working methodology draft covering implementation Steps 1-14"
    properties.author = ""
    properties.keywords = "market microstructure, Q-learning, replication, methodology"

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(10)
    kicker.paragraph_format.space_after = Pt(5)
    kicker_run = kicker.add_run("METHODOLOGY WORKING DRAFT")
    set_font(kicker_run, "Calibri", 9.5)
    kicker_run.bold = True
    kicker_run.font.color.rgb = BLUE

    doc.add_paragraph(
        "Methodology Working Draft: Market Design, Benchmarks, and Finite States",
        style="Title",
    )
    doc.add_paragraph(
        "Replication of Dou, Goldstein, and Ji (2025) | Implementation Steps 1-14",
        style="Subtitle",
    )

    metadata = doc.add_paragraph()
    metadata.paragraph_format.space_before = Pt(0)
    metadata.paragraph_format.space_after = Pt(14)
    metadata_run = metadata.add_run(
        "Prepared 28 August 2026  |  Python replication  |  Living document"
    )
    set_font(metadata_run, "Calibri", 9.5)
    metadata_run.font.color.rgb = GRAY

    add_draft_note(
        doc,
        "Scope note.",
        "This draft documents the static market foundation, theoretical pricing rule, "
        "Nash and perfect-cartel strategies, benchmark profits, finite action and price "
        "grids, and the finite state representation. Hand-checkable examples and "
        "exhaustive index checks are software validations. They are not empirical "
        "estimates or evidence that the full learning model has been replicated.",
    )

    doc.add_heading("1. Replication Design and Current Scope", level=1)
    add_body_paragraph(
        doc,
        "The replication follows the economic environment and numerical setup in "
        "Dou, Goldstein, and Ji (2025). Development is deliberately incremental: each "
        "economic relationship is first implemented as a separate Python component, "
        "validated against a deterministic calculation, and only then connected to "
        "other components. This design makes the mapping from the paper to the code "
        "explicit and reduces the risk that an error in one equation is concealed by "
        "a larger simulation.",
    )
    add_body_paragraph(
        doc,
        "The present implementation covers the fundamental-value process, noise "
        "trading, aggregate order flow, information-insensitive demand, informed-"
        "speculator profits, the market maker's objective, the theoretical pricing rule, "
        "the Nash and perfect-cartel fixed points and expected profits, the finite action "
        "and price grids, and the state-index system. Q-learning agents, the adaptive "
        "market maker, convergence criteria, and repeated-session experiments remain "
        "outside the current scope and will be added only after their corresponding code "
        "has been validated.",
    )
    add_body_paragraph(
        doc,
        "The baseline parameter values relevant to the components implemented so far "
        "are I = 2 informed speculators, an unconditional fundamental mean of v-bar = 1, "
        "fundamental volatility sigma_v = 1, n_v = 10 value-grid points, information-"
        "insensitive demand slope xi = 500, and pricing-error weight theta = 0.1. The "
        "finite-grid settings are n_x = 15 actions, n_p = 31 prices, and widening "
        "parameter iota = 0.1. The paper considers both a low-noise environment with "
        "sigma_u = 0.1 and a high-noise environment with sigma_u = 100.",
    )

    doc.add_heading("2. Economic Environment", level=1)
    doc.add_heading("2.1 Participants and within-period information", level=2)
    add_body_paragraph(
        doc,
        "Time is discrete. In each period, a short-lived asset is traded and its "
        "fundamental value is realized at the end of that period. The market contains "
        "I risk-neutral informed speculators, a representative noise trader, a "
        "representative information-insensitive investor sector, and a representative "
        "market maker. Exogenous shocks are independent and identically distributed "
        "across periods.",
    )
    add_body_paragraph(
        doc,
        "Each informed speculator observes the current fundamental value perfectly before "
        "submitting an order, but does not observe the contemporaneous noise order. The "
        "market maker subsequently observes aggregate informed-plus-noise order flow, not "
        "its separate components. The information-insensitive sector follows a known "
        "price-dependent demand schedule. These information restrictions are central to "
        "the price-formation mechanism.",
    )

    doc.add_heading("2.2 Fundamental value and Gaussian discretization", level=2)
    add_body_paragraph(
        doc,
        "The period-t fundamental value follows the normal distribution assumed in the "
        "paper:",
    )
    add_equation(doc, "vₜ ~ N(v̄, σᵥ²).")
    add_body_paragraph(
        doc,
        "The baseline normalization is v-bar = 1 and sigma_v = 1. Because the later "
        "learning model requires a finite state space, the continuous distribution is "
        "approximated using n_v = 10 equiprobable grid values. For k = 1, ..., n_v, the "
        "probability midpoint and corresponding grid point are:",
    )
    add_equation(doc, "qₖ = (2k - 1)/(2nᵥ),      vₖ = v̄ + σᵥ Φ⁻¹(qₖ).")
    add_body_paragraph(
        doc,
        "The midpoint probabilities are 0.05, 0.15, ..., 0.95, and each resulting value "
        "receives probability 1/n_v. Probability midpoints avoid zero and one, where the "
        "Gaussian quantile function is unbounded. The resulting grid is ordered and "
        "symmetric around one. Its arithmetic mean is 1.000000 and its discrete standard "
        "deviation is computed as:",
    )
    add_equation(doc, "σ̂ᵥ = [(1/nᵥ) Σₖ(vₖ - v̄)²]¹ᐟ² = 0.937970 ≈ 0.938.")
    add_body_paragraph(
        doc,
        "This discretization is a numerical approximation to the paper's assumed normal "
        "distribution. It is not an empirical estimate of the distribution of real asset "
        "fundamentals. Both normality and the baseline parameter values are inherited "
        "modeling assumptions.",
    )

    doc.add_heading("2.3 Noise trading", level=2)
    add_body_paragraph(
        doc,
        "The representative noise trader submits an order that is independent of the "
        "fundamental value and other exogenous shocks:",
    )
    add_equation(doc, "uₜ ~ N(0, σᵤ²).")
    add_body_paragraph(
        doc,
        "The paper studies sigma_u = 0.1 and sigma_u = 100. Moving from the former to "
        "the latter multiplies the standard deviation by 1,000 and the variance by "
        "1,000,000. It does not change the population mean, which remains zero. A finite "
        "sample mean may nevertheless be nonzero, and scaling a fixed sequence of "
        "standard-normal draws scales that finite-sample mean together with the sample "
        "standard deviation.",
    )
    add_body_paragraph(
        doc,
        "For preliminary reproducibility checks, Gaussian orders were generated with a "
        "fixed random seed of 42. Two generators initialized with the same seed reproduced "
        "the same first draw. Distributional behavior was checked using 100,000 draws, "
        "requiring the sample mean to be near zero and the population-form sample standard "
        "deviation to lie within 2 percent of the specified sigma_u. A final policy for "
        "deriving independent session-level seeds will be specified before the full "
        "simulation stage. The seed, sample size, and tolerance are implementation "
        "diagnostics rather than choices prescribed by the paper.",
    )

    doc.add_heading("2.4 Aggregate order flow and observability", level=2)
    add_body_paragraph(
        doc,
        "Speculator i submits order x_i,t. Aggregate informed orders and the noise order "
        "form the total order flow:",
    )
    add_equation(doc, "yₜ = Σᵢ xᵢ,ₜ + uₜ.")
    add_body_paragraph(
        doc,
        "With two informed speculators, y_t = x_1,t + x_2,t + u_t. The market maker "
        "observes y_t only after these orders have been submitted. It cannot determine "
        "how much of the observed total came from informed trading and how much came from "
        "the noise trader. Noise trading is therefore unobservable but not price-"
        "irrelevant: a realization of u_t changes y_t and can affect price both directly "
        "and through the market maker's inference about fundamental value.",
    )

    doc.add_heading("2.5 Information-insensitive investor demand", level=2)
    add_body_paragraph(
        doc,
        "The representative information-insensitive sector responds deterministically to "
        "the deviation of the market price from the unconditional expected fundamental "
        "value. Its aggregate demand is equation (3.2) of the source paper:",
    )
    add_equation(doc, "zₜ = -ξ(pₜ - v̄),      ξ ≥ 0.")
    add_body_paragraph(
        doc,
        "Demand is downward sloping. The sector sells when p_t exceeds v-bar, buys when "
        "p_t is below v-bar, and submits no order when the two are equal. This participant "
        "is economically distinct from the noise trader: u_t is random and enters y_t "
        "before price formation, whereas z_t is a deterministic function of the price "
        "schedule known to the market maker.",
    )

    doc.add_heading("2.6 Informed-speculator payoff", level=2)
    add_body_paragraph(
        doc,
        "The period payoff of informed speculator i is the value-price difference "
        "multiplied by its signed position:",
    )
    add_equation(doc, "πᵢ,ₜ = (vₜ - pₜ)xᵢ,ₜ.")
    add_body_paragraph(
        doc,
        "A positive order denotes a purchase and a negative order denotes a sale or short "
        "position. Buying is profitable when fundamental value exceeds price; shorting is "
        "profitable when price exceeds fundamental value. This period payoff will later "
        "serve as the immediate reward in each informed agent's Q-learning update.",
    )

    doc.add_heading("3. Market-Maker Pricing", level=1)
    doc.add_heading("3.1 Objective function", level=2)
    add_body_paragraph(
        doc,
        "After observing y_t, the market maker chooses p_t to minimize the conditional "
        "expected loss specified in equation (3.3) of the paper:",
    )
    add_equation(doc, "minₚ E[(yₜ + zₜ)² + θ(pₜ - vₜ)² | yₜ].")
    add_body_paragraph(
        doc,
        "The first term is a quadratic inventory cost. The market maker must absorb the "
        "position -(y_t + z_t), so a larger residual imbalance is more costly. The second "
        "term penalizes squared pricing error, and theta > 0 controls the weight placed on "
        "price accuracy. A lower objective value is preferable; the objective is a loss "
        "measure rather than a trading profit.",
    )
    add_body_paragraph(
        doc,
        "The expectation is conditional on observed aggregate order flow. Because z_t is "
        "deterministic for a proposed price, uncertainty enters through v_t. In the "
        "standalone objective test, the expectation was evaluated over a deliberately "
        "simple conditional distribution so that the result could be verified by hand. "
        "That artificial distribution is not a calibrated belief used in the final model.",
    )

    doc.add_heading("3.2 Theoretical price rule", level=2)
    add_body_paragraph(
        doc,
        "Substituting z_t = -xi(p_t - v-bar) into the objective and differentiating with "
        "respect to p_t produces the first-order condition:",
    )
    add_equation(doc, "(ξ² + θ)pₜ = ξyₜ + ξ²v̄ + θE[vₜ | yₜ].")
    add_body_paragraph(
        doc,
        "Solving this condition yields equation (3.4) of the paper:",
    )
    add_equation(
        doc,
        "pₜ = [ξ/(ξ² + θ)]yₜ + [ξ²/(ξ² + θ)]v̄ "
        "+ [θ/(ξ² + θ)]E[vₜ | yₜ].",
    )
    add_body_paragraph(
        doc,
        "The three components capture the direct price response to aggregate order flow, "
        "the weight placed on the unconditional mean fundamental value, and the weight "
        "placed on value inferred from order flow. The conditional expectation is the "
        "market maker's belief after observing y_t; it is not the realized fundamental "
        "value and need not coincide with one of the discrete grid points.",
    )

    doc.add_heading("3.3 How noise trading affects price", level=2)
    add_body_paragraph(
        doc,
        "The pricing rule does not take u_t as a separate observable input because the "
        "market maker cannot identify it. Nevertheless, substituting y_t = sum_i x_i,t + "
        "u_t shows that noise affects price through two channels. First, it changes the "
        "order imbalance in the direct y_t term. Second, it changes the signal from which "
        "the market maker forms E[v_t | y_t]. The noise variance therefore governs how "
        "informative a given order-flow realization is, even though the realized noise "
        "order remains hidden.",
    )
    add_body_paragraph(
        doc,
        "The standalone Step 7 pricing function accepts E[v_t | y_t] directly so that "
        "equation (3.4) can be tested in isolation. Steps 8-10 then determine this "
        "conditional expectation analytically from the benchmark relationship between "
        "value and order flow, including sigma_u. The later adaptive market maker will "
        "instead estimate the relationship from a rolling historical sample.",
    )

    doc.add_heading("4. Analytical Benchmark Strategies and Profits", level=1)
    doc.add_heading("4.1 Non-collusive Nash benchmark", level=2)
    add_body_paragraph(
        doc,
        "The non-collusive Nash benchmark is the one-period equilibrium in which each "
        "informed speculator maximizes its own expected profit while treating the other "
        "speculators' equilibrium orders as given. The benchmark uses the linear pricing "
        "rule and order-flow decomposition stated in Online Appendix Proposition IA.1:",
    )
    add_equation(
        doc,
        "pᴺ(yₜ) = v̄ + λᴺyₜ,      yₜ = xᵢ,ₜ + (I - 1)xᴺ(vₜ) + uₜ.",
    )
    add_body_paragraph(
        doc,
        "The symmetric informed strategy and its trading-intensity coefficient are:",
    )
    add_equation(
        doc,
        "xᴺ(vₜ) = χᴺ(vₜ - v̄),      χᴺ = 1/[(I + 1)λᴺ].",
    )
    add_body_paragraph(
        doc,
        "A positive value signal produces a purchase, whereas a negative signal produces "
        "a short position. The factor I + 1 reflects strategic price impact: an informed "
        "speculator recognizes that increasing its own order raises the transaction price "
        "and reduces the order's profitability. The benchmark is non-collusive because a "
        "single trader may deviate while the remaining traders' orders stay fixed.",
    )

    doc.add_heading("4.2 Perfect-cartel benchmark", level=2)
    add_body_paragraph(
        doc,
        "The perfect-cartel benchmark is the opposite strategic extreme. All informed "
        "speculators coordinate their orders and maximize joint expected profit as if "
        "they were a single monopolistic trader. Online Appendix Proposition IA.2 uses:",
    )
    add_equation(
        doc,
        "pᴹ(yₜ) = v̄ + λᴹyₜ,      yₜ = Ixᴹ(vₜ) + uₜ.",
    )
    add_body_paragraph(
        doc,
        "The common order assigned to each cartel member and its intensity are:",
    )
    add_equation(
        doc,
        "xᴹ(vₜ) = χᴹ(vₜ - v̄),      χᴹ = 1/(2Iλᴹ).",
    )
    add_body_paragraph(
        doc,
        "Unlike the Nash calculation, the cartel changes all members' orders together "
        "and internalizes the price effect of their combined flow. It therefore restricts "
        "informed trading relative to the competitive benchmark. This perfect-cartel "
        "solution is a theoretical reference point; it is not yet a claim that the later "
        "independent Q-learning agents collude.",
    )

    doc.add_heading("4.3 Coupled benchmark fixed points", level=2)
    add_body_paragraph(
        doc,
        "Steps 8 and 9 initially treated lambda^N and lambda^M as transparent unit-test "
        "inputs. Step 10 closes each benchmark by jointly determining trading intensity, "
        "the market maker's inference from aggregate order flow, and equilibrium price "
        "impact. For benchmark B in {N, M}, the inference slope is:",
    )
    add_equation(
        doc,
        "E[vₜ | yₜ] = v̄ + γᴮyₜ,      γᴮ = Iχᴮ / [(Iχᴮ)² + (σᵤ/σ̂ᵥ)²].",
    )
    add_body_paragraph(
        doc,
        "The corresponding price-impact coefficient is:",
    )
    add_equation(
        doc,
        "λᴮ = (θγᴮ + ξ)/(θ + ξ²).",
    )
    add_body_paragraph(
        doc,
        "The implementation uses sigma_v_hat = 0.937969795249 from the ten-point "
        "fundamental-value grid, rather than the nominal continuous sigma_v = 1. Noise "
        "trading enters through sigma_u/sigma_v_hat and changes how informative observed "
        "order flow is. Information-insensitive investors enter through xi and influence "
        "how strongly price must respond to order flow. The complete Nash system combines "
        "this pair of equations with chi^N = 1/[(I + 1)lambda^N]; the cartel system "
        "instead uses chi^M = 1/(2I lambda^M). Here gamma is the analytical benchmark "
        "projection coefficient, not the rolling regression estimate used by the later "
        "adaptive market maker.",
    )

    doc.add_heading("4.4 Numerical solution method", level=2)
    add_body_paragraph(
        doc,
        "The source paper specifies the coupled fixed-point equations but does not name a "
        "numerical root-finding algorithm. Bisection is therefore an implementation choice "
        "in this replication, not an additional economic assumption and not a procedure "
        "attributed to the authors. For each benchmark, the code evaluates the residual:",
    )
    add_equation(
        doc,
        "Rᴮ(λ) = λ - [θγᴮ(χᴮ(λ)) + ξ]/(θ + ξ²).",
    )
    add_body_paragraph(
        doc,
        "The economically admissible solution satisfies Rᴮ(λ) = 0 and λ > 0. "
        "For the strictly positive parameters used here, the residual has exactly one "
        "positive root. Since gamma cannot exceed sigma_v_hat/(2sigma_u), that root lies "
        "inside the analytical interval:",
    )
    add_equation(
        doc,
        "ξ/(θ + ξ²) ≤ λᴮ ≤ [ξ + θσ̂ᵥ/(2σᵤ)]/(θ + ξ²).",
    )
    add_body_paragraph(
        doc,
        "Bisection repeatedly halves this bracket and retains the half containing the "
        "residual's zero. Iteration stops when the absolute residual or the interval's "
        "half-width is no greater than 10⁻¹⁸, subject to a maximum of 200 iterations. "
        "The solver is designed to run once per distinct parameter configuration before "
        "the repeated market simulation; the later experiment runner will store and reuse "
        "the result rather than recomputing it every trading period. A local timing "
        "diagnostic completed "
        "10,000 solves in approximately 0.40 seconds after import, or about 40 microseconds "
        "per root. This timing is an implementation diagnostic, not a result reported by "
        "the paper.",
    )

    doc.add_heading("4.5 Expected profits under the analytical benchmarks", level=2)
    add_body_paragraph(
        doc,
        "After the fixed points have been solved, Step 11 calculates the unconditional "
        "expected profit of one informed speculator in each theoretical benchmark. The "
        "Nash expression and the perfect-cartel expression are:",
    )
    add_equation(
        doc,
        "πᴺ = σ̂ᵥ²/[(I + 1)²λᴺ],      πᴹ = σ̂ᵥ²/(4Iλᴹ).",
    )
    add_body_paragraph(
        doc,
        "The second quantity is the expected profit of one cartel member, not the joint "
        "profit of all cartel members. Joint cartel profit is Iπᴹ. These quantities are "
        "reference values for the later learning experiment: πᴺ is the non-collusive "
        "benchmark and πᴹ is the perfectly coordinated benchmark. Neither value is "
        "evidence that the later independent learning agents have colluded.",
    )
    add_body_paragraph(
        doc,
        "Using I = 2 and σ̂ᵥ² = 0.879787336800, the low-noise calibration gives πᴺ = "
        "48.877064491239 and πᴹ = 54.986686555306 per trader. The corresponding high-"
        "noise values are 48.877067212825 and 54.986693324281. The formulas were checked "
        "independently by rebuilding x = χ(v - v̄), y = Ix + u, p = v̄ + λy, and "
        "π = (v - p)x over all ten value points and the symmetric shocks "
        "u in {-σᵤ, +σᵤ}. Formula and direct calculation agreed to within "
        "7.11 × 10⁻¹⁵. The symmetric pair is a deterministic mean-zero test device, not "
        "a replacement for Gaussian noise in the final simulation.",
    )

    doc.add_page_break()
    doc.add_heading("5. Finite Action, Price, and State Spaces", level=1)
    add_body_paragraph(
        doc,
        "Q-learning requires a finite action set and a finite set of state labels. Steps "
        "12-14 therefore convert the analytical benchmark ranges into 15 action choices, "
        "31 lagged-price labels, and 3,100 observable states. These grids do not replace "
        "the continuous economic calculations. They define what a learning agent can "
        "choose and how its information is stored.",
    )

    doc.add_heading("5.1 Action grid", level=2)
    add_body_paragraph(
        doc,
        "The paper specifies nₓ = 15 permitted actions and the widening parameter "
        "ι = 0.1. The implementation first builds a reusable multiplier interval from "
        "the perfect-cartel and Nash trading intensities:",
    )
    add_equation(
        doc,
        "dχ = χᴺ - χᴹ,      cᴸ = χᴹ - ιdχ,      cᴴ = χᴺ + ιdχ.",
    )
    add_equation(
        doc,
        "cⱼ = cᴸ + [j/(nₓ - 1)](cᴴ - cᴸ),      j = 0, ..., nₓ - 1.",
    )
    add_body_paragraph(
        doc,
        "For current fundamental value v, multiplier cⱼ becomes the signed order",
    )
    add_equation(doc, "xⱼ(v) = (v - v̄)cⱼ.")
    add_body_paragraph(
        doc,
        "The 15 actions are therefore not 15 globally fixed quantities. Each of the ten "
        "fundamental values has its own row of 15 raw orders, giving a 10 × 15 value-"
        "action array. If v exceeds v̄, the orders are purchases; if v is below v̄, the "
        "negative value signal produces mirrored short-sale quantities. The multiplier "
        "representation does not force the learned policy to be linear: the Q-learning "
        "agent may choose a different action index in every state.",
    )
    add_body_paragraph(
        doc,
        "Under low noise, the widened multiplier interval is "
        "[120.833281667, 170.833301667] with spacing 3.571429999999. Under high noise, "
        "it is [120.833297665, 170.833310336] with spacing 3.571429476510. Validation "
        "confirmed exactly 15 ordered choices in each environment, correct endpoints and "
        "spacing, and equal-magnitude opposite-signed orders at symmetric fundamental "
        "values. The grid is deterministic once the parameters are fixed and uses no "
        "random seed.",
    )

    doc.add_heading("5.2 Price grid", level=2)
    add_body_paragraph(
        doc,
        "The market maker will produce a continuous transaction price, but the lagged "
        "price inside the learning state requires a finite label. The paper constructs "
        "initial bounds from the most negative and most positive Nash or cartel orders "
        "over the value grid:",
    )
    add_equation(
        doc,
        "pᴸ = v̄ + λᴺ[I min(xᴹ, xᴺ) - 1.96σᵤ],",
    )
    add_equation(
        doc,
        "pᴴ = v̄ + λᴺ[I max(xᴹ, xᴺ) + 1.96σᵤ].",
    )
    add_body_paragraph(
        doc,
        "The interval is widened by ι on both sides and divided into nₚ = 31 equally "
        "spaced points:",
    )
    add_equation(
        doc,
        "p̃ᴸ = pᴸ - ι(pᴴ - pᴸ),      p̃ᴴ = pᴴ + ι(pᴴ - pᴸ),",
    )
    add_equation(
        doc,
        "Pₘ = p̃ᴸ + [m/(nₚ - 1)](p̃ᴴ - p̃ᴸ),      m = 0, ..., nₚ - 1.",
    )
    add_body_paragraph(
        doc,
        "The code reproduces the coefficient 1.96 printed in the paper's formula. "
        "Statistically, plus or minus 1.96 standard deviations describes the central "
        "95 percent normal interval, whose endpoints are approximately the 2.5th and "
        "97.5th percentiles. Nearby prose in the paper instead calls them the 5th and "
        "95th percentiles. Because the numerical formula is explicit, this replication "
        "uses 1.96 exactly and records the wording discrepancy rather than silently "
        "changing the coefficient.",
    )
    add_grid_summary_table(doc)
    add_body_paragraph(
        doc,
        "Both grids are symmetric around v̄ = 1, and index 15 is exactly 1. The high-"
        "noise grid is wider because its bounds explicitly contain 1.96σᵤ. Validation "
        "confirmed 31 ordered, equally spaced points and exact agreement with the widened "
        "bounds. Constructing P does not constrain the market maker to quote only one of "
        "these prices; P is used to label the lagged price in the learning state.",
    )

    doc.add_heading("5.3 State information and the continuous-price mapping", level=2)
    add_body_paragraph(
        doc,
        "The paper defines an informed agent's state as",
    )
    add_equation(doc, "sₜ = (pₜ₋₁, vₜ₋₁, vₜ) ∈ P × V × V.")
    add_body_paragraph(
        doc,
        "The state contains the previous transaction price, the previous fundamental "
        "value, and the current fundamental value. It does not contain the current price "
        "pₜ because traders must choose their current orders before pₜ is determined. "
        "After trading and the next value draw, the next state will be "
        "sₜ₊₁ = (pₜ, vₜ, vₜ₊₁).",
    )
    add_body_paragraph(
        doc,
        "Fundamental values are drawn directly from the ten-point grid V, so vₜ₋₁ and "
        "vₜ must match permitted values. The implementation rejects an off-grid value "
        "instead of silently assigning it to a different economic state. The transaction "
        "price is different: it is continuous, whereas the state requires an element of "
        "the finite grid P.",
    )
    add_draft_note(
        doc,
        "Explicit replication choice.",
        "The paper and online appendix define P but do not state how a continuous realized "
        "price should be mapped to it. This implementation selects the nearest point in P, "
        "clips observations below or above the grid to the corresponding endpoint, and "
        "selects the lower point at an exact midpoint. The full simulation will record the "
        "frequency of endpoint clipping.",
    )
    add_body_paragraph(
        doc,
        "Only the state label is discretized. The original continuous price remains the "
        "transaction price used to calculate information-insensitive demand zₜ, informed-"
        "speculator profits, and the observations stored in the market maker's historical "
        "sample. Mapping the lagged state price therefore does not replace or round the "
        "price used by the economic model.",
    )

    doc.add_heading("5.4 State indexing and exhaustive validation", level=2)
    add_body_paragraph(
        doc,
        "With nₚ = 31 and nᵥ = 10, the finite state space contains",
    )
    add_equation(doc, "|S| = nₚnᵥ² = 31 × 10 × 10 = 3,100 states.")
    add_body_paragraph(
        doc,
        "Each state is first stored as three zero-based indexes: previous-price index jₚ, "
        "previous-value index jᵥ,lag, and current-value index jᵥ. For convenient later "
        "array access, the tuple is reversibly compressed into one integer:",
    )
    add_equation(
        doc,
        "jₛ = (jₚnᵥ + jᵥ,lag)nᵥ + jᵥ.",
    )
    add_body_paragraph(
        doc,
        "This numbering has no economic meaning; it is only a storage convention. It "
        "assigns the 3,100 states the integers 0 through 3,099. For example, in the low-"
        "noise environment, a continuous previous price of 1.10 maps to "
        "P[16] = 1.087756887. If the previous and current fundamentals are "
        "V[2] = 0.325510250 and V[7] = 1.674489750, the state tuple is (16, 2, 7) "
        "and its integer identifier is (16 × 10 + 2) × 10 + 7 = 1,627.",
    )
    add_body_paragraph(
        doc,
        "Validation mapped every one of the ten value points and every point in both "
        "31-point price grids back to its original index. It tested observations on both "
        "sides of a midpoint, the exact-midpoint convention, and clipping below and above "
        "the price grid. Finally, all 3,100 state tuples were enumerated: every tuple had "
        "a unique integer identifier, and decoding every identifier reproduced the "
        "original tuple. An arbitrary off-grid fundamental value was correctly rejected. "
        "No Q-table or learning rule is introduced in this step; a later Q-table will "
        "have 3,100 × 15 = 46,500 state-action cells per agent.",
    )

    doc.add_heading("6. Implementation and Validation Evidence", level=1)
    doc.add_heading("6.1 Modular implementation", level=2)
    add_body_paragraph(
        doc,
        "The prototype is implemented in Python 3.13. Each validated relationship is "
        "exposed as a small reusable function. Standalone scripts retain their own fixed "
        "examples behind a main-entry guard, allowing them to run independently while "
        "also being imported into integration checks without rerunning demonstration "
        "code. Assertions compare computed results with hand-derived values and stop "
        "execution if a discrepancy is detected.",
    )

    doc.add_heading("6.2 Component-level checks", level=2)
    add_body_paragraph(
        doc,
        "The fundamental-value routine was checked for ten values, the intended midpoint "
        "probabilities, symmetry, a mean of one, and a discrete standard deviation of "
        "approximately 0.938. The noise routine was checked for seed reproducibility and "
        "distributional moments. Aggregate order flow was verified using x_1,t = 2, "
        "x_2,t = -1, and u_t = 0.5, which gives y_t = 1.5.",
    )
    add_body_paragraph(
        doc,
        "Using the baseline v-bar = 1 and xi = 500, prices of 1.01, 0.99, and 1.00 "
        "generated information-insensitive orders of -5, 5, and 0, respectively. The "
        "payoff function was checked across buying and short-selling positions under both "
        "underpricing and overpricing. An order magnitude of two and a value-price gap of "
        "0.20 produced profits or losses of 0.40 with the expected signs.",
    )

    doc.add_heading("6.3 Integration Checkpoint A", level=2)
    add_body_paragraph(
        doc,
        "The first integration checkpoint connected aggregate order flow, information-"
        "insensitive demand, and both informed traders' payoffs in one deterministic "
        "example. With v_t = 1.20, a temporarily fixed p_t = 1.01, x_1,t = 2, x_2,t = -1, "
        "and u_t = 0.5, the connected functions returned y_t = 1.5, z_t = -5, and "
        "y_t + z_t = -3.5. The two informed traders earned 0.38 and -0.19. The nonzero "
        "residual imbalance is expected because the checkpoint deliberately fixed the "
        "price before the market-maker pricing rule was connected.",
    )

    doc.add_heading("6.4 Market-maker and price-rule checks", level=2)
    add_body_paragraph(
        doc,
        "The objective evaluator was first checked with y_t = 10, v-bar = 1, xi = 500, "
        "theta = 0.1, and a two-point conditional belief assigning equal weight to 0.8 and "
        "1.2. Candidate prices 1.00, 1.02, and 1.04 generated objective values 100.00400, "
        "0.00404, and 100.00416. The comparison demonstrates that the coded objective "
        "rewards the price that nearly eliminates the residual imbalance. It does not "
        "restrict the paper's market maker to three candidate prices.",
    )
    add_body_paragraph(
        doc,
        "The theoretical price function was then verified independently using y_t = 0.5, "
        "v-bar = 1, E[v_t | y_t] = 1.2, xi = 2, and theta = 1. The three terms in the "
        "pricing equation were 0.20, 0.80, and 0.24, giving p_t = 1.24. The Step 6 "
        "objective was lower at 1.24 than at nearby prices 1.14 and 1.34. These small "
        "numbers were chosen for hand calculation and are not the paper's baseline "
        "calibration.",
    )
    add_body_paragraph(
        doc,
        "Two limiting cases provide additional checks. When xi = 0, information-"
        "insensitive demand is inactive and the formula reduces to p_t = E[v_t | y_t]. "
        "When xi is extremely large, the price approaches v-bar + y_t/xi, the value that "
        "induces information-insensitive demand to offset the observed flow. For v-bar = "
        "1, y_t = 0.5, and xi = 1,000,000, this clearing limit is 1.0000005. The values "
        "1.2, 1.24, and 1.0000005 are diagnostic outputs, not substantive results.",
    )

    doc.add_heading("6.5 Nash and cartel strategy checks", level=2)
    add_body_paragraph(
        doc,
        "The Step 8 Nash calculation was first isolated using I = 2, lambda^N = 0.5, "
        "v-bar = 1, and v_t = 1.3. These deliberately non-calibrated inputs imply "
        "chi^N = 2/3 and an order of 0.20 per trader. With E[u_t] = 0, aggregate flow "
        "is 0.40 and expected price is 1.20. Holding the other trader's order fixed, "
        "own orders of 0.10, 0.20, and 0.30 yield expected profits of 0.015, 0.020, "
        "and 0.015. The central order is therefore the individual best response in the "
        "test, and (I + 1)lambda^N chi^N - 1 has a zero residual.",
    )
    add_body_paragraph(
        doc,
        "The Step 9 cartel check used the same I, v-bar, v_t, and test price impact, "
        "but varied both members' orders together. It gives chi^M = 0.5 and an order of "
        "0.15 per member. With expected noise zero, total flow is 0.30, expected price "
        "is 1.15, profit per member is 0.0225, and joint profit is 0.045. Coordinated "
        "orders of 0.10 and 0.20 each yield the lower joint profit 0.040, while "
        "2I lambda^M chi^M - 1 has a zero residual. These numbers validate the formula; "
        "they are not the paper-calibrated benchmark coefficients.",
    )
    add_body_paragraph(
        doc,
        "Expected noise is set to zero only when checking the analytical first-order "
        "conditions. Realized-noise tests retain the noise trader. In the Nash example, "
        "u_t = 0.1 raises flow from 0.40 to 0.50 and price from 1.20 to 1.25. In the "
        "cartel example, it raises flow from 0.30 to 0.40 and price from 1.15 to 1.20, "
        "reducing profit per member from 0.0225 to 0.015.",
    )

    doc.add_page_break()
    doc.add_heading("6.6 Fixed-point checks and baseline outputs", level=2)
    add_body_paragraph(
        doc,
        "The coupled equations were solved with I = 2, xi = 500, theta = 0.1, and "
        "sigma_v_hat = 0.937969795249. Both paper noise environments were evaluated. "
        "The results below are numerical solutions implied by the paper's equations and "
        "calibration; the paper does not print these coefficient values directly.",
    )
    add_fixed_point_results_table(doc)
    add_body_paragraph(
        doc,
        "Across all four solutions, the absolute fixed-point residual was no larger "
        "than 8.7 × 10⁻¹⁹, and the Nash or cartel first-order-condition identity residual "
        "was no larger than 1.2 × 10⁻¹⁶. These diagnostics establish numerical consistency "
        "with the implemented equations; they do not establish empirical validity or "
        "completion of the full replication.",
    )
    add_body_paragraph(
        doc,
        "The cartel intensity is lower than the Nash intensity in both environments, "
        "consistent with coordinated restriction of informed order flow. All four lambda "
        "values remain close to 1/xi = 0.002 because theta/(theta + xi^2) is approximately "
        "4 × 10⁻⁷, so the inference term receives little weight in the price-impact "
        "coefficient. This proximity does not make noise irrelevant: changing sigma_u "
        "materially changes gamma, and the standard deviation of the directly noise-driven "
        "price component is approximately lambda sigma_u, about 0.0002 under low noise "
        "and 0.20 under high noise.",
    )

    doc.add_heading("7. Current Boundary of the Replication", level=1)
    add_body_paragraph(
        doc,
        "The completed checks establish internal consistency for the static market "
        "foundation, theoretical pricing rule, Nash and perfect-cartel strategies, their "
        "coupled fixed points and expected profits, the 15-action grid, the 31-point price "
        "grid, and the 3,100-state representation. They do not establish that the complete "
        "implementation reproduces the paper's quantitative findings or that independent "
        "AI agents have learned to collude. The immediate next stage is Step 15: draw and "
        "validate the initial state uniformly from P × V × V. Later stages will initialize "
        "the Q-tables, implement action selection and Q updates for two independent agents, "
        "add the adaptive market maker's rolling regressions, assemble the exact within-"
        "period protocol, and reproduce the convergence criteria, outcome measures, "
        "mechanism tests, and comparative-static experiments.",
    )
    add_body_paragraph(
        doc,
        "This working methodology will be revised after each validated implementation "
        "checkpoint. Final claims about replication success will be made only after the "
        "full simulation has been run with saved configurations, reproducible random "
        "streams, and explicit comparison tolerances against the paper's reported "
        "results.",
    )

    doc.add_heading("References", level=1)
    reference = doc.add_paragraph(style="Reference")
    reference.add_run(
        "Dou, W. W., Goldstein, I., and Ji, Y. (2025). AI-Powered Trading, "
        "Algorithmic Collusion, and Price Efficiency. NBER Working Paper No. 34054. "
        "National Bureau of Economic Research. https://www.nber.org/papers/w34054"
    )

    return doc


def main() -> None:
    """Build and save the DOCX file."""

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
